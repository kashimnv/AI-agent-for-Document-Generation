"""
Assembles the final .docx deliverable from the plan + generated section
content, using python-docx. Produces a polished, consistently styled
document: title page, generated-on date, section headings, intro
paragraphs, bullet lists, and (for a couple of doc types) a simple table.
"""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)  # dark blue


def _style_title_page(doc: Document, title: str, doc_type: str, user_request: str):
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = ACCENT_COLOR

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run(doc_type.replace("_", " ").title())
    sub_run.font.size = Pt(14)
    sub_run.italic = True

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')} by Autonomous Document Agent")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    req_heading = doc.add_paragraph()
    req_run = req_heading.add_run("Original Request")
    req_run.bold = True
    req_run.font.size = Pt(11)

    req_body = doc.add_paragraph(user_request)
    req_body.runs[0].italic = True

    doc.add_page_break()


def _add_section(doc: Document, heading: str, content: dict, level: int = 1):
    doc.add_heading(heading, level=level)

    paragraph_text = content.get("paragraph", "")
    if paragraph_text:
        doc.add_paragraph(paragraph_text)

    bullets = content.get("bullets", [])
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_paragraph()  # spacer


def _add_task_trace_appendix(doc: Document, task_trace: list):
    doc.add_page_break()
    doc.add_heading("Appendix: Agent Execution Log", level=1)
    doc.add_paragraph(
        "The following steps were autonomously planned and executed by the "
        "agent to produce this document:"
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Step"
    hdr_cells[1].text = "Action"
    hdr_cells[2].text = "Status"

    for task in task_trace:
        row_cells = table.add_row().cells
        row_cells[0].text = str(task.get("step", ""))
        row_cells[1].text = str(task.get("action", ""))
        row_cells[2].text = str(task.get("status", "")).upper()


def build_document(
    output_path: Path,
    title: str,
    doc_type: str,
    user_request: str,
    sections: list,
    section_content: dict,
    task_trace: list,
) -> Path:
    doc = Document()

    # Base font
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    _style_title_page(doc, title, doc_type, user_request)

    doc.add_heading("Table of Contents (Sections)", level=1)
    for section_name in sections:
        doc.add_paragraph(section_name, style="List Number")
    doc.add_page_break()

    for section_name in sections:
        content = section_content.get(section_name, {})
        _add_section(doc, section_name, content)

    _add_task_trace_appendix(doc, task_trace)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
